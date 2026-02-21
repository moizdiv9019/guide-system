from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

def generate_roadmap_docx(roadmap_data):
    doc = Document()

    title = doc.add_heading('Personalized Career Roadmap', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for month_data in roadmap_data:
        doc.add_heading(
            f"Month {month_data['month']}: {month_data['focus']}", 
            level=1
        )

        for week in month_data['weeks']:
            doc.add_heading(
                f"Week {week['week_number']}: {week['week_title']}", 
                level=2
            )

            doc.add_paragraph("Learning Objectives:", style='List Bullet')

            for obj in week['learning_objectives']:
                p = doc.add_paragraph(obj, style='List Bullet 2')
                p.paragraph_format.left_indent = Pt(40)

            tools_str = ", ".join(week['tools'])
            doc.add_paragraph(f"Tools: {tools_str}")

            task_para = doc.add_paragraph()
            run = task_para.add_run("Practice Task: ")
            run.bold = True
            task_para.add_run(week['practice_task'])

            doc.add_paragraph()

    # Save to memory instead of file
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
